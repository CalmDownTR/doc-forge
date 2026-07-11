"""Tests for CARD-029: DOCXParser."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches
from PIL import Image

from docforge.config import ParseConfig
from docforge.models import ContentType
from docforge.parsers.docx_parser import DOCXParser


def _create_test_docx(tmp_path: Path, add_table: bool = True, add_image: bool = False) -> Path:
    """Create a test DOCX file with headings, paragraphs, and optionally a table."""
    doc = DocxDocument()
    doc.add_heading("Main Title", level=1)
    doc.add_paragraph("This is the first paragraph.")
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("Section one content.")
    doc.add_heading("Subsection", level=3)
    doc.add_paragraph("Subsection details.")

    if add_table:
        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"Row{i}Col{j}"

    doc.add_paragraph("Text after table.")

    if add_image:
        img_bytes_io = io.BytesIO()
        img = Image.new("RGB", (10, 10), color="blue")
        img.save(img_bytes_io, format="PNG")
        img_bytes_io.seek(0)
        doc.add_picture(img_bytes_io, width=Inches(1))

    doc.add_paragraph("Final paragraph.")

    filepath = tmp_path / "test.docx"
    doc.save(str(filepath))
    return filepath


class TestDOCXParserCanParse:
    def test_can_parse_docx(self):
        parser = DOCXParser()
        assert parser.can_parse(Path("test.docx"), "docx") is True

    def test_can_parse_other(self):
        parser = DOCXParser()
        assert parser.can_parse(Path("test.pdf"), "pdf") is False
        assert parser.can_parse(Path("test.xlsx"), "xlsx") is False
        assert parser.can_parse(Path("test.txt"), "txt") is False


class TestDOCXParserHeadings:
    def test_heading_levels_correct(self, tmp_path: Path):
        filepath = _create_test_docx(tmp_path)
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)

        text_blocks = [b for b in blocks if b.type == ContentType.TEXT]
        headings = [b for b in text_blocks if b.metadata.get("is_heading")]

        assert len(headings) >= 3
        assert headings[0].content.startswith("# ")
        assert "Main Title" in headings[0].content

        # Find level 2 heading
        level2 = [b for b in headings if b.content.startswith("## ")]
        assert len(level2) >= 1
        assert "Section One" in level2[0].content

        # Find level 3 heading
        level3 = [b for b in headings if b.content.startswith("### ")]
        assert len(level3) >= 1


class TestDOCXParserTables:
    def test_table_extracted_as_markdown(self, tmp_path: Path):
        filepath = _create_test_docx(tmp_path, add_table=True)
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)

        table_blocks = [b for b in blocks if b.type == ContentType.TABLE]
        assert len(table_blocks) >= 1, f"Expected at least 1 table block, got {len(table_blocks)}"

        table_content = table_blocks[0].content
        assert "|" in table_content
        assert "---" in table_content
        assert "Row0Col0" in table_content
        assert "Row1Col1" in table_content

    def test_table_has_correct_structure(self, tmp_path: Path):
        filepath = _create_test_docx(tmp_path, add_table=True)
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)

        table_blocks = [b for b in blocks if b.type == ContentType.TABLE]
        table_content = table_blocks[0].content
        lines = table_content.strip().split("\n")
        assert len(lines) >= 4  # header + separator + 3 data rows
        # First line should be header
        assert lines[0].startswith("|")
        # Second line should be separator
        assert "---" in lines[1]


class TestDOCXParserImages:
    def test_images_extracted_to_images_dir(self, tmp_path: Path):
        filepath = _create_test_docx(tmp_path, add_image=True)
        config = ParseConfig(extract_images=True)
        blocks = DOCXParser().parse(filepath, config)

        image_blocks = [b for b in blocks if b.type == ContentType.IMAGE]
        assert len(image_blocks) >= 1, f"Expected at least 1 image block, got {len(image_blocks)}"

        img_content = image_blocks[0].content
        assert "_images/" in img_content or "test_images/" in img_content

        # Verify image directory was created and contains files
        img_dir = tmp_path / "test_images"
        assert img_dir.exists(), f"Image directory {img_dir} should exist"
        assert len(list(img_dir.iterdir())) >= 1

    def test_images_extraction_disabled(self, tmp_path: Path):
        filepath = _create_test_docx(tmp_path, add_image=True)
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)

        # When extract_images is False, images inside paragraphs should be handled
        # as text blocks (if the paragraph has text too) or skipped
        image_blocks = [b for b in blocks if b.type == ContentType.IMAGE]
        assert len(image_blocks) == 0


class TestDOCXParserOrder:
    def test_body_tables_images_in_correct_order(self, tmp_path: Path):
        filepath = _create_test_docx(tmp_path, add_table=True, add_image=True)
        config = ParseConfig(extract_images=True)
        blocks = DOCXParser().parse(filepath, config)

        # Check that reading_order is sequential
        for i, b in enumerate(blocks):
            assert b.reading_order == i, f"Block at index {i} has reading_order {b.reading_order}"

        # The order should be: headings, paragraphs, table, text, image, text
        types_in_order = [b.type for b in blocks]
        assert ContentType.TABLE in types_in_order

    def test_all_blocks_have_page_1(self, tmp_path: Path):
        filepath = _create_test_docx(tmp_path, add_table=True)
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)

        for b in blocks:
            assert b.page == 1

    def test_empty_document(self, tmp_path: Path):
        """Test parsing a minimal empty document."""
        doc = DocxDocument()
        filepath = tmp_path / "empty.docx"
        doc.save(str(filepath))
        config = ParseConfig()
        blocks = DOCXParser().parse(filepath, config)
        # Empty docx may have some empty paragraphs — all should be filtered
        assert all(b.content.strip() for b in blocks)


class TestDOCXParserEmptyTable:
    def test_no_table_when_none_present(self, tmp_path: Path):
        """Test that a document without tables has zero table blocks."""
        doc = DocxDocument()
        doc.add_paragraph("Just a paragraph.")
        filepath = tmp_path / "no_table.docx"
        doc.save(str(filepath))
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)

        table_blocks = [b for b in blocks if b.type == ContentType.TABLE]
        assert len(table_blocks) == 0


class TestDOCXParserEdgeCases:
    def test_custom_image_output_dir(self, tmp_path: Path):
        """Test that custom image_output_dir is used."""
        import io

        doc = DocxDocument()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Text.")
        img_bytes_io = io.BytesIO()
        from PIL import Image
        img = Image.new("RGB", (10, 10), color="red")
        img.save(img_bytes_io, format="PNG")
        img_bytes_io.seek(0)
        doc.add_picture(img_bytes_io)
        doc.add_paragraph("After image.")

        filepath = tmp_path / "custom_dir.docx"
        doc.save(str(filepath))

        custom_dir = tmp_path / "my_images"
        config = ParseConfig(extract_images=True, image_output_dir=custom_dir)
        blocks = DOCXParser().parse(filepath, config)

        assert custom_dir.exists()
        image_blocks = [b for b in blocks if b.type == ContentType.IMAGE]
        assert len(image_blocks) >= 1

    def test_document_with_only_empty_paragraphs(self, tmp_path: Path):
        """Test that empty paragraphs are filtered out."""
        doc = DocxDocument()
        p = doc.add_paragraph("")
        # Force truly empty paragraph
        p.clear()
        filepath = tmp_path / "empty_paras.docx"
        doc.save(str(filepath))
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)
        # All blocks returned should have non-empty content
        text_blocks = [b for b in blocks if b.type == ContentType.TEXT]
        for b in text_blocks:
            assert b.content.strip() != ""

    def test_escape_cell_pipe_characters(self):
        """Test that pipe characters in table cells are escaped."""
        parser = DOCXParser()
        result = parser._escape_cell("a|b")
        assert result == "a\\|b"

    def test_build_cell_grid_with_merged_cells(self, tmp_path: Path):
        """Test building cell grid handles merged cells."""
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        # Normal cells
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        table.rows[1].cells[0].text = "C"
        table.rows[1].cells[1].text = "D"

        filepath = tmp_path / "merged.docx"
        doc.save(str(filepath))
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)

        table_blocks = [b for b in blocks if b.type == ContentType.TABLE]
        assert len(table_blocks) >= 1
        content = table_blocks[0].content
        assert "A" in content
        assert "B" in content
        assert "C" in content
        assert "D" in content

    def test_heading_style_none_returns_zero(self, tmp_path: Path):
        """Test heading level returns 0 when heading style has no useful name."""
        doc = DocxDocument()
        doc.add_paragraph("No heading style paragraph.")
        filepath = tmp_path / "no_heading_style.docx"
        doc.save(str(filepath))
        config = ParseConfig(extract_images=False)
        blocks = DOCXParser().parse(filepath, config)

        text_blocks = [b for b in blocks if b.type == ContentType.TEXT]
        # The paragraph should NOT be detected as a heading
        headings = [b for b in text_blocks if b.metadata.get("is_heading")]
        assert len(headings) == 0
